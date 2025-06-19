import os
import PyPDF2
import re
from docx import Document
import win32com.client as win32
import uuid  # Para gerar nomes aleatórios

# Caminho da pasta onde estão os PDFs
pasta_pdf = r'C:\Users\Visitante\Desktop\exportar DNG\mesclados'
pasta_temp = os.path.join(pasta_pdf, 'temp')  # Definindo o caminho da pasta 'temp'

# Garantir que a pasta 'temp' exista
if not os.path.exists(pasta_temp):
    os.makedirs(pasta_temp)

caminho_word = r'C:\Users\Visitante\Desktop\exportar DNG\imagem\VOLUME ZERADO.docx'

# Função para extrair o texto de um PDF
def extrair_texto_pdf(caminho_pdf):
    with open(caminho_pdf, 'rb') as arquivo:
        leitor_pdf = PyPDF2.PdfReader(arquivo)
        texto = ""
        for pagina in range(len(leitor_pdf.pages)):
            texto += leitor_pdf.pages[pagina].extract_text()
        return texto

# Função para procurar "VOLUME" e verificar o valor anterior
def verificar_volume_zerado(texto_pdf):
    resultado = re.search(r'(\d+)\s*VOLUME', texto_pdf)
    if resultado:
        valor_antes_volume = resultado.group(1)  # Pega o número antes da palavra "VOLUME"
        if valor_antes_volume == "0":  # Verifica se o valor é 0
            return True
    return False

# Função para substituir os campos no Word e gerar PDF
def substituir_e_gerar_pdf(numero_nf):
    doc = Document(caminho_word)
    
    # Substituindo o número 654321 pelo número da NF
    for paragrafo in doc.paragraphs:
        if '654321' in paragrafo.text:
            paragrafo.text = paragrafo.text.replace('654321', str(numero_nf))
    
    # Substituindo o número 123456 pelo número da NF - 1
    numero_nf_anterior = numero_nf - 1
    for paragrafo in doc.paragraphs:
        if '123456' in paragrafo.text:
            paragrafo.text = paragrafo.text.replace('123456', str(numero_nf_anterior))
    
    # Gerar um nome aleatório para o arquivo
    nome_aleatorio = str(uuid.uuid4())  # Gerar um nome único para o arquivo

    # Caminho dos arquivos temporários na pasta 'temp'
    caminho_docx_temp = os.path.join(pasta_temp, f'{nome_aleatorio}.docx')
    caminho_pdf_temp = os.path.join(pasta_temp, f'{nome_aleatorio}.pdf')
    
    # Salvando o documento alterado
    doc.save(caminho_docx_temp)

    # Convertendo o Word para PDF
    word = win32.Dispatch("Word.Application")
    doc = word.Documents.Open(caminho_docx_temp)
    doc.SaveAs(caminho_pdf_temp, FileFormat=17)  # 17 é o formato PDF
    doc.Close()
    word.Quit()

    return caminho_pdf_temp

# Função para mesclar PDFs
def mesclar_pdfs(pdf_lista, caminho_pdf_saida):
    pdf_writer = PyPDF2.PdfWriter()
    
    for pdf in pdf_lista:
        with open(pdf, 'rb') as arquivo_pdf:
            pdf_reader = PyPDF2.PdfReader(arquivo_pdf)
            for pagina in range(len(pdf_reader.pages)):
                pdf_writer.add_page(pdf_reader.pages[pagina])

    with open(caminho_pdf_saida, 'wb') as arquivo_saida:
        pdf_writer.write(arquivo_saida)

# Função para encontrar e processar as NFs zeradas
def encontrar_nfs_zeradas():
    arquivos_zerados = []
    for arquivo in os.listdir(pasta_pdf):
        if arquivo.endswith('.pdf'):
            caminho_pdf = os.path.join(pasta_pdf, arquivo)
            texto_pdf = extrair_texto_pdf(caminho_pdf)
            if verificar_volume_zerado(texto_pdf):
                numero_nf = int(arquivo.split('.')[0])  # Considera o nome do arquivo como o número da NF
                arquivos_zerados.append((numero_nf, caminho_pdf))
    return arquivos_zerados

# Função para renomear arquivos antes de excluir
def renomear_arquivos(originais):
    renomeados = []
    for arquivo in originais:
        novo_nome = arquivo.replace('.pdf', ' - original.pdf')
        novo_caminho = os.path.join(pasta_pdf, os.path.basename(novo_nome))
        os.rename(arquivo, novo_caminho)
        renomeados.append(novo_caminho)
    return renomeados

# Função para excluir arquivos PDF
def excluir_arquivos(*arquivos):
    for arquivo in arquivos:
        if os.path.exists(arquivo):
            os.remove(arquivo)
            print(f"Arquivo {arquivo} excluído.")

# Chama a função e processa as NFs zeradas
nfs_zeradas = encontrar_nfs_zeradas()

if nfs_zeradas:
    for numero_nf, caminho_pdf_zerada in nfs_zeradas:
        print(f"Processando a NF zerada {numero_nf}...")
        
        # Determina o caminho do arquivo PDF da NF anterior
        numero_nf_anterior = numero_nf - 1
        caminho_pdf_anterior = os.path.join(pasta_pdf, f"{numero_nf_anterior}.pdf")

        # Renomeia os arquivos antes de excluir
        arquivos_renomeados = renomear_arquivos([caminho_pdf_zerada, caminho_pdf_anterior])
        
        # Gerar PDF a partir do Word
        caminho_pdf_aviso = substituir_e_gerar_pdf(numero_nf)
        
        # Mesclar os PDFs: aviso, NF anterior e NF zerada
        caminho_pdf_saida = os.path.join(pasta_pdf, f"{numero_nf}.pdf")
        mesclar_pdfs([caminho_pdf_aviso] + arquivos_renomeados, caminho_pdf_saida)
        
        print(f"Arquivo gerado: {caminho_pdf_saida}")
        
        # Excluir os arquivos renomeados
        excluir_arquivos(*arquivos_renomeados)
else:
    print("Não foram encontradas NFs com volume zerado.")
